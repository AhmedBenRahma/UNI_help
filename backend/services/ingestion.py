"""Document ingestion pipeline for UniHelp."""

from __future__ import annotations

import hashlib
from datetime import datetime, timezone
from pathlib import Path

from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from pypdf import PdfReader

from backend.core.config import Settings
from backend.core.exceptions import IngestionError
from backend.core.logging import get_logger
from backend.services.vector_store import VectorStoreService


class IngestionService:
    """Service to read, chunk, and index documents."""

    def __init__(self, settings: Settings, vector_store: VectorStoreService) -> None:
        self.settings = settings
        self.vector_store = vector_store
        self.logger = get_logger(__name__)
        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.settings.CHUNK_SIZE,
            chunk_overlap=self.settings.CHUNK_OVERLAP,
        )

    def ingest_file(self, file_path: Path) -> dict[str, int | list[str]]:
        """Ingest one file and persist new chunks to vector store."""
        self.logger.info("ingestion_started", filename=file_path.name)
        if not file_path.exists():
            raise IngestionError(f"File not found: {file_path}")

        raw_documents = self._load_file(file_path)
        if not raw_documents:
            raise IngestionError(f"No readable content in file: {file_path.name}")

        content_hash = self._compute_hash(raw_documents)
        existing_hashes = self.vector_store.existing_hashes()
        if content_hash in existing_hashes:
            self.logger.info("ingestion_deduplicated", filename=file_path.name, content_hash=content_hash)
            return {"indexed_chunks": 0, "source_files": [file_path.name]}

        timestamp = datetime.now(timezone.utc).isoformat()
        enriched_docs = self._attach_metadata(raw_documents, file_path.name, content_hash, timestamp)
        chunks = self.splitter.split_documents(enriched_docs)
        self.vector_store.add_documents(chunks)

        self.logger.info("ingestion_completed", filename=file_path.name, chunks_indexed=len(chunks))
        return {"indexed_chunks": len(chunks), "source_files": [file_path.name]}

    def _load_file(self, file_path: Path) -> list[Document]:
        """Load file into LangChain documents with page metadata."""
        suffix = file_path.suffix.lower()
        try:
            if suffix == ".pdf":
                return self._load_pdf(file_path)
            if suffix == ".docx":
                return self._load_docx(file_path)
            if suffix == ".txt":
                return self._load_txt(file_path)
            raise IngestionError(f"Unsupported file type: {suffix}")
        except (OSError, ValueError) as exc:
            self.logger.error("ingestion_load_failed", filename=file_path.name, error=str(exc))
            raise IngestionError(f"Failed to load file {file_path.name}") from exc

    def _load_pdf(self, file_path: Path) -> list[Document]:
        """Load PDF by pages."""
        reader = PdfReader(str(file_path))
        documents: list[Document] = []
        for idx, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                documents.append(Document(page_content=text, metadata={"page_number": idx}))
        return documents

    def _load_docx(self, file_path: Path) -> list[Document]:
        """Load DOCX as a single page-equivalent document."""
        doc = DocxDocument(str(file_path))
        text = "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())
        if not text.strip():
            return []
        return [Document(page_content=text, metadata={"page_number": 1})]

    def _load_txt(self, file_path: Path) -> list[Document]:
        """Load TXT as a single page-equivalent document."""
        text = file_path.read_text(encoding="utf-8")
        if not text.strip():
            return []
        return [Document(page_content=text, metadata={"page_number": 1})]

    def _compute_hash(self, documents: list[Document]) -> str:
        """Compute MD5 hash of full normalized content."""
        concatenated = "\n".join(doc.page_content.strip() for doc in documents if doc.page_content.strip())
        return hashlib.md5(concatenated.encode("utf-8"), usedforsecurity=False).hexdigest()

    def _attach_metadata(
        self,
        documents: list[Document],
        filename: str,
        content_hash: str,
        upload_timestamp: str,
    ) -> list[Document]:
        """Attach standard metadata to loaded documents."""
        enriched: list[Document] = []
        for doc in documents:
            metadata = {
                **doc.metadata,
                "filename": filename,
                "upload_timestamp": upload_timestamp,
                "content_hash": content_hash,
            }
            enriched.append(Document(page_content=doc.page_content, metadata=metadata))
        return enriched
